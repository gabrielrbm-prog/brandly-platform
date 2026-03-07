from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_bold_paragraph(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_check_item(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('[ ] ').bold = True
    p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    return table

# ===================== CONTENT =====================

# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GUIA COMPLETO')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Como Abrir sua Loja no TikTok Shop Brasil')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Produto: Cafe Soluvel')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Preparado por: Brandly')
run.font.size = Pt(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Data: Marco/2026')
run.font.size = Pt(12)

doc.add_page_break()

# SOBRE O TIKTOK SHOP
add_heading_styled('SOBRE O TIKTOK SHOP BRASIL', 1)
doc.add_paragraph(
    'O TikTok Shop e a funcionalidade de e-commerce dentro do TikTok que permite vender produtos '
    'diretamente na plataforma. Lancado no Brasil em maio de 2025, ja conta com milhares de vendedores '
    'ativos. O grande diferencial e que o consumidor compra sem sair do app, direto de videos, lives e '
    'da vitrine do perfil.'
)
p = doc.add_paragraph()
run = p.add_run('Beneficio para novos vendedores: 0% de comissao da plataforma nos primeiros 90 dias.')
run.bold = True

# NOTA ALIMENTOS
add_heading_styled('NOTA SOBRE A CATEGORIA DE ALIMENTOS', 2)
doc.add_paragraph(
    'A categoria de Alimentos e Bebidas no TikTok Shop Brasil foi liberada apos o lancamento inicial. '
    'Marcas como BIS (Mondelez) ja operam oficialmente na plataforma vendendo alimentos. Para vender '
    'cafe soluvel, sera necessario solicitar habilitacao da categoria e apresentar documentacao sanitaria. '
    'Este guia detalha tudo o que voce precisa.'
)

doc.add_page_break()

# FASE 1
add_heading_styled('FASE 1 — PRE-REQUISITOS', 1)
doc.add_paragraph('Antes de iniciar o cadastro, certifique-se de ter tudo pronto:')

add_heading_styled('1.1 Documentacao Empresarial', 2)
add_table(
    ['Documento', 'Detalhes'],
    [
        ['CNPJ ativo', 'MEI serve, mas empresa completa (ME, LTDA) da mais credibilidade e acesso a recursos'],
        ['CNAE compativel', 'Deve incluir atividade de venda de alimentos/bebidas'],
        ['Contrato Social ou CCMEI', 'Dependendo do tipo de empresa'],
        ['Documento do representante legal', 'RG, CNH ou passaporte (frente e verso)'],
        ['Conta bancaria PJ', 'No nome do CNPJ, para recebimento das vendas'],
        ['Comprovante de endereco comercial', 'Pode ser exigido na verificacao'],
    ]
)

doc.add_paragraph()
add_heading_styled('1.2 Documentacao Especifica para Alimentos', 2)
add_table(
    ['Documento', 'Detalhes'],
    [
        ['Registro ou notificacao ANVISA', 'Cafe soluvel embalado precisa de registro ou protocolo de isencao'],
        ['Alvara sanitario', 'Emitido pela vigilancia sanitaria municipal ou estadual'],
        ['Rotulo em conformidade', 'Conforme RDC 429/2020 — tabela nutricional, ingredientes, lote, validade, dados do fabricante'],
        ['Laudo laboratorial', 'Analise do produto (pode ser solicitado pelo TikTok)'],
        ['Certificado de Boas Praticas', 'Se voce for o fabricante do cafe'],
    ]
)

doc.add_paragraph()
add_heading_styled('1.3 Logistica e Estoque', 2)
items_logistica = [
    'O estoque DEVE estar em armazem localizado no Brasil (envio internacional nao e aceito)',
    'Definir transportadora ou Correios para envio',
    'Embalagem adequada para alimentos (protecao, conservacao, lacre de seguranca)',
    'Prazo de entrega impacta diretamente o ranking da sua loja — quanto mais rapido, melhor',
]
for item in items_logistica:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# FASE 2
add_heading_styled('FASE 2 — CADASTRO NO TIKTOK SHOP (PASSO A PASSO)', 1)

add_heading_styled('Passo 1: Criar ou ajustar sua conta no TikTok', 2)
steps = [
    'Baixe o app TikTok no celular (se ainda nao tiver)',
    'Crie uma conta ou use uma conta existente',
    'Converta para Conta Comercial: Va em Perfil > Menu > Configuracoes > Conta > "Mudar para Conta Comercial" (gratuito)',
    'Escolha a categoria do seu negocio',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

add_heading_styled('Passo 2: Acessar o Seller Center (Central do Vendedor)', 2)
steps = [
    'No computador, acesse: https://seller-br.tiktok.com/account/register',
    'Faca login com sua conta TikTok ou cadastre-se com e-mail/telefone comercial',
    'Aceite os termos de servico',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

add_heading_styled('Passo 3: Selecionar tipo de vendedor', 2)
doc.add_paragraph('Voce tera duas opcoes:')
add_bold_paragraph('Opcao A — Corporacao (LTDA, SA, EIRELI)')
for s in ['Upload do contrato social da empresa', 'Upload do documento do representante legal (frente e verso)', 'Informacoes da empresa (razao social, CNPJ, endereco)']:
    doc.add_paragraph(s, style='List Bullet')
add_bold_paragraph('Opcao B — Empresa Individual / MEI')
for s in ['Upload do CCMEI (Certificado de Condicao de Microempreendedor Individual)', 'Upload do documento pessoal (frente e verso)']:
    doc.add_paragraph(s, style='List Bullet')

add_heading_styled('Passo 4: Definir nome e informacoes da loja', 2)
for s in [
    'Escolha um nome para a loja (ex: "Cafe [Sua Marca]", "[Marca] Oficial")',
    'Preencha o endereco comercial completo',
    'Selecione a categoria principal: Alimentos e Bebidas',
    'Se a categoria nao aparecer, sera necessario solicitar habilitacao (ver Fase 3)',
]:
    doc.add_paragraph(s, style='List Bullet')

add_heading_styled('Passo 5: Vincular conta bancaria', 2)
for s in ['Nome da conta (deve ser o nome da empresa/CNPJ)', 'Banco', 'Numero da agencia', 'Numero da conta', 'A conta DEVE ser PJ — nao e aceito conta de pessoa fisica']:
    doc.add_paragraph(s, style='List Bullet')

add_heading_styled('Passo 6: Enviar para analise', 2)
for s in [
    'Revise todas as informacoes e clique em "Enviar"',
    'O TikTok analisa o cadastro em 1 a 6 dias uteis',
    'O resultado sera enviado por e-mail',
    'Se recusado, voce recebera o motivo e podera corrigir e reenviar',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# FASE 3
add_heading_styled('FASE 3 — HABILITACAO DA CATEGORIA DE ALIMENTOS', 1)
doc.add_paragraph('Apos a aprovacao do cadastro basico, voce precisa liberar a categoria de alimentos:')

add_heading_styled('Passo 7: Acessar o Qualification Center', 2)
for s in [
    'Dentro do Seller Center, va em "Qualification Center" (Centro de Qualificacao)',
    'Clique em "Add Category Authorization" (Adicionar Autorizacao de Categoria)',
    'Selecione: Alimentos e Bebidas > Cafe / Bebidas em Po (ou categoria mais proxima)',
]:
    doc.add_paragraph(s, style='List Number')

add_heading_styled('Passo 8: Enviar documentacao sanitaria', 2)
doc.add_paragraph('Faca upload dos seguintes documentos:')
for s in [
    'Registro ou protocolo de notificacao ANVISA',
    'Alvara sanitario vigente',
    'Foto do rotulo completo do produto (todas as faces da embalagem)',
    'Laudo tecnico ou certificado de qualidade',
    'Nota fiscal de compra ou fabricacao (comprovando a origem do produto)',
]:
    doc.add_paragraph(s, style='List Bullet')

add_heading_styled('Passo 9: Aguardar aprovacao', 2)
for s in [
    'A analise da categoria pode levar alguns dias adicionais',
    'Se negado, verifique qual documento faltou ou esta incorreto e reenvie',
    'Em caso de duvida, acione o suporte do TikTok Shop via chat no Seller Center',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# FASE 4
add_heading_styled('FASE 4 — CADASTRO DO PRODUTO', 1)

add_heading_styled('Passo 10: Adicionar o cafe soluvel na loja', 2)
doc.add_paragraph('No Seller Center, va em Products > Add Product e preencha:')

add_bold_paragraph('Titulo do produto')
doc.add_paragraph('Descritivo e com palavras-chave. Exemplos:')
for s in [
    '"Cafe Soluvel Premium 200g - 100% Arabica - [Sua Marca]"',
    '"Cafe Soluvel Gourmet [Marca] - Torra Media - 100g"',
]:
    doc.add_paragraph(s, style='List Bullet')

add_bold_paragraph('Descricao')
for s in [
    'Escreva em portugues, de forma detalhada e atrativa',
    'Inclua: origem do cafe, tipo de grao, nivel de torra, modo de preparo, diferenciais, peso liquido',
    'Destaque certificacoes (organico, sustentavel, etc.) se houver',
]:
    doc.add_paragraph(s, style='List Bullet')

add_bold_paragraph('Imagens (minimo 3, ideal 5+)')
for s in [
    'Foto principal do produto (fundo branco ou lifestyle)',
    'Foto do produto aberto / po do cafe',
    'Foto de uma xicara servida',
    'Foto do rotulo / informacoes nutricionais',
    'Foto de ambientacao / estilo de vida',
]:
    doc.add_paragraph(s, style='List Number')

add_bold_paragraph('Video do produto')
doc.add_paragraph('Altamente recomendado — mostre o produto, o preparo e alguem degustando.')

add_bold_paragraph('Preco, estoque e variacoes')
for s in [
    'Defina um preco competitivo (pesquise concorrentes)',
    'Insira a quantidade em estoque',
    'Configure variacoes se tiver (ex: 100g, 200g, 500g)',
    'Preencha peso e dimensoes para calculo de frete',
]:
    doc.add_paragraph(s, style='List Bullet')

add_heading_styled('Passo 11: Configurar frete e politica de devolucao', 2)
for s in [
    'Defina prazo de envio (recomendado: 1-3 dias uteis para despacho)',
    'Configure a politica de devolucao (obrigatorio para ativar a loja)',
    'Defina as regioes de entrega (todo Brasil ou regioes especificas)',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# FASE 5
add_heading_styled('FASE 5 — ESTRATEGIA DE VENDAS', 1)

add_heading_styled('Passo 12: Criar conteudo de venda', 2)
p = doc.add_paragraph()
run = p.add_run('No TikTok Shop, sem video nao ha venda. Conteudo e o motor das vendas.')
run.bold = True

for s in [
    'Videos curtos (15-60s): Preparo do cafe, unboxing, degustacao, comparacao, receitas com cafe',
    'Etiqueta de produto: Em todo video, adicione a tag laranja do produto (carrinho) para compra direta',
    'Lives de venda: Live commerce e o formato que MAIS converte no TikTok Shop',
    'Frequencia: Poste pelo menos 1 video por dia com a tag do produto',
]:
    doc.add_paragraph(s, style='List Bullet')

add_heading_styled('Passo 13: Ativar o programa de afiliados', 2)
doc.add_paragraph('Uma das formas mais poderosas de vender no TikTok Shop:')
for s in [
    'No Seller Center, va em "Affiliate"',
    'Crie um plano de comissao (ex: 10-20% por venda)',
    'Creators do TikTok poderao adicionar seu cafe na vitrine deles e divulgar',
    'Voce so paga comissao quando a venda acontece',
]:
    doc.add_paragraph(s, style='List Number')

add_heading_styled('Passo 14: Aproveitar o periodo de 0% de comissao', 2)
doc.add_paragraph('Nos primeiros 90 dias, o TikTok nao cobra comissao. Use esse periodo para:')
for s in [
    'Gerar o maximo de vendas e avaliacoes positivas',
    'Testar precos e ofertas',
    'Construir reputacao da loja',
]:
    doc.add_paragraph(s, style='List Bullet')

add_heading_styled('Passo 15: Investir em anuncios (opcional, mas recomendado)', 2)
for s in [
    'Dentro do Seller Center, crie campanhas de Product Shopping Ads',
    'Comece com orcamento baixo (R$20-50/dia) e escale o que converter',
    'Os anuncios aparecem no feed dos usuarios com o botao de compra direto',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# TAXAS
add_heading_styled('RESUMO DE TAXAS E CUSTOS', 1)
add_table(
    ['Item', 'Valor'],
    [
        ['Cadastro na plataforma', 'Gratuito'],
        ['Comissao por venda (primeiros 90 dias)', '0%'],
        ['Comissao por venda (apos 90 dias)', '~5% (varia por categoria)'],
        ['Comissao de afiliados', 'Definida por voce (ex: 10-20%)'],
        ['Anuncios (opcional)', 'A partir de R$20/dia'],
    ]
)

doc.add_paragraph()

# CHECKLIST
add_heading_styled('CHECKLIST — IMPRIMA E VA MARCANDO', 1)
checklist_items = [
    'CNPJ ativo com CNAE de alimentos',
    'Alvara sanitario em dia',
    'Registro ou notificacao ANVISA do cafe soluvel',
    'Rotulo em conformidade com a legislacao (RDC 429/2020)',
    'Conta no TikTok convertida para Conta Comercial',
    'Cadastro no Seller Center (seller-br.tiktok.com) — enviado e aprovado',
    'Documentos da empresa enviados e verificados',
    'Conta bancaria PJ vinculada',
    'Categoria de Alimentos habilitada no Qualification Center',
    'Produto cadastrado com fotos, descricao e preco',
    'Frete e politica de devolucao configurados',
    'Programa de afiliados ativado com comissao definida',
    'Primeiro video publicado com tag do produto',
    'Primeira live de vendas realizada',
]
for item in checklist_items:
    add_check_item(item)

doc.add_paragraph()

# LINKS
add_heading_styled('LINKS UTEIS', 1)
links = [
    ('Seller Center Brasil (cadastro)', 'https://seller-br.tiktok.com/account/register'),
    ('TikTok Shop Academy (tutoriais)', 'Disponivel dentro do Seller Center apos login'),
    ('Suporte TikTok Shop', 'Chat disponivel dentro do Seller Center'),
    ('Guia oficial de configuracao', 'https://ads.tiktok.com/help/article/set-up-tiktok-shop-using-tiktok-seller-center?lang=pt'),
]
for name, url in links:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(url)

doc.add_paragraph()

# DICAS FINAIS
add_heading_styled('DICAS FINAIS', 1)
dicas = [
    ('Conteudo e rei', 'No TikTok Shop, quem vende mais e quem produz mais conteudo. Invista em videos criativos e lives frequentes.'),
    ('Afiliados aceleram', 'Ative o programa de afiliados desde o primeiro dia. Creators divulgando seu cafe = vendas no piloto automatico.'),
    ('Avaliacoes importam', 'Incentive clientes satisfeitos a avaliarem o produto. Lojas com boas avaliacoes ganham mais destaque.'),
    ('Responda rapido', 'O tempo de resposta a mensagens de clientes impacta a reputacao da loja.'),
    ('Acompanhe metricas', 'Use o dashboard do Seller Center para monitorar vendas, visualizacoes, conversao e ajustar sua estrategia.'),
]
for i, (title, desc) in enumerate(dicas, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title} — ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Documento preparado com base em informacoes oficiais do TikTok Shop Brasil e fontes publicas '
    'atualizadas ate marco/2026. Recomendamos sempre verificar as politicas vigentes diretamente no '
    'Seller Center, pois o TikTok pode atualizar regras e categorias a qualquer momento.'
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Save
output_path = '/Users/gabrielrubim/Documents/Plataformas/SaaS/Brandly/Guia_TikTok_Shop_Cafe_Soluvel.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
